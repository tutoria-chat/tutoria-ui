#!/usr/bin/env python3
"""
Build the TutorIA user guide in 3 languages, from ONE content source.

For each locale it writes:
  - public/guides/user-guide-<locale>.docx   (styled Word document)
  - public/guides/user-guide-<locale>.json   (blocks the in-app viewer renders)

Run:  python scripts/build_user_guide.py
Requires: python-docx  (pip install python-docx)

Content lives in content.py (CONTENT dict). Block tuples:
  ('h2', id, text) ('h3', text) ('p', text)
  ('ul', [items]) ('ol', [items])
  ('tip'|'warn'|'info', text)
"""
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

from content import CONTENT

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.normpath(os.path.join(HERE, "..", "public", "guides"))

BRAND = RGBColor(0x5E, 0x17, 0xEB)      # TutorIA purple
ACCENT = RGBColor(0x5C, 0xE1, 0xE6)     # cyan
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x66, 0x66, 0x66)

CALLOUT_COLORS = {
    "tip": ("E8FBFC", "0E7C86", "💡"),
    "warn": ("FFF4E5", "9A5B00", "⚠️"),
    "info": ("F1ECFE", "4A12B8", "ℹ️"),
}


def _shade(el, hexfill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexfill)
    el.append(shd)


def _add_callout(doc, variant, text):
    fill, ink, emoji = CALLOUT_COLORS[variant]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _shade(cell._tc.get_or_add_tcPr(), fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{emoji}  {text}")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(ink)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _toc_field(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldStart = OxmlElement("w:fldChar"); fldStart.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    hint = OxmlElement("w:t"); hint.text = "Right-click → Update Field to build the table of contents."
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    r = run._r
    for node in (fldStart, instr, fldSep, hint, fldEnd):
        r.append(node)


def build_docx(locale, data, path):
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    for lvl, color, size in (("Heading 1", BRAND, 18), ("Heading 2", BRAND, 14), ("Heading 3", ACCENT, 12)):
        st = doc.styles[lvl]
        st.font.color.rgb = color if lvl != "Heading 3" else RGBColor(0x0E, 0x7C, 0x86)
        st.font.size = Pt(size)
        st.font.bold = True

    # Cover
    doc.add_paragraph().paragraph_format.space_after = Pt(60)
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title.add_run("TutorIA"); trun.font.size = Pt(48); trun.font.bold = True; trun.font.color.rgb = BRAND
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run(data["title"]); srun.font.size = Pt(20); srun.font.color.rgb = INK
    desc = doc.add_paragraph(); desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    drun = desc.add_run(data["subtitle"]); drun.font.size = Pt(12); drun.font.italic = True; drun.font.color.rgb = MUTED
    upd = doc.add_paragraph(); upd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    urun = upd.add_run(data["updated"]); urun.font.size = Pt(10); urun.font.color.rgb = MUTED
    doc.add_page_break()

    # TOC
    h = doc.add_paragraph(data["tocTitle"], style="Heading 1")
    _toc_field(doc)
    doc.add_page_break()

    for _id, heading, blocks in data["sections"]:
        doc.add_paragraph(heading, style="Heading 2")
        for block in blocks:
            kind = block[0]
            if kind == "h3":
                doc.add_paragraph(block[1], style="Heading 3")
            elif kind == "p":
                doc.add_paragraph(block[1])
            elif kind == "ul":
                for item in block[1]:
                    doc.add_paragraph(item, style="List Bullet")
            elif kind == "ol":
                for item in block[1]:
                    doc.add_paragraph(item, style="List Number")
            elif kind in CALLOUT_COLORS:
                _add_callout(doc, kind, block[1])

    doc.save(path)


def build_json(locale, data, path):
    sections = []
    for _id, heading, blocks in data["sections"]:
        out_blocks = []
        for block in blocks:
            kind = block[0]
            if kind in ("h3", "p", "tip", "warn", "info"):
                out_blocks.append({"type": kind, "text": block[1]})
            elif kind in ("ul", "ol"):
                out_blocks.append({"type": kind, "items": block[1]})
        sections.append({"id": _id, "heading": heading, "blocks": out_blocks})
    payload = {
        "title": data["title"],
        "subtitle": data["subtitle"],
        "updated": data["updated"],
        "tocTitle": data["tocTitle"],
        "sections": sections,
    }
    with open(path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=1)


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    for locale, data in CONTENT.items():
        build_docx(locale, data, os.path.join(OUT_DIR, f"user-guide-{locale}.docx"))
        build_json(locale, data, os.path.join(OUT_DIR, f"user-guide-{locale}.json"))
        print(f"  [ok] {locale}: docx + json")
    print(f"Done -> {OUT_DIR}")


if __name__ == "__main__":
    main()
